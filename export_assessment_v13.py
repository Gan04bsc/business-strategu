from dataclasses import dataclass
from pathlib import Path
import re
import sys
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = Path(r'd:\study\Business Strategy')


def resolve_path(arg, default):
    if arg is None:
        return default
    p = Path(arg)
    return p if p.is_absolute() else BASE / p


SRC = resolve_path(sys.argv[1] if len(sys.argv) > 1 else None, BASE / 'Assessment_new_v13_working.md')
OUT = resolve_path(sys.argv[2] if len(sys.argv) > 2 else None, BASE / 'Assessment_new_v13.docx')


@dataclass
class ReferenceRecord:
    raw_text: str
    clean_text: str
    bookmark: str
    author_block: str
    year: str
    author_forms: list[str]


def sanitize_bookmark(text: str, index: int) -> str:
    cleaned = re.sub(r'[^A-Za-z0-9_]+', '_', text).strip('_')
    if not cleaned:
        cleaned = f'ref_{index}'
    if cleaned[0].isdigit():
        cleaned = f'ref_{cleaned}'
    return f'{cleaned[:28]}_{index}'


def normalize(text: str) -> str:
    return re.sub(r'\s+', ' ', text.strip()).lower()


def parse_blocks(text: str):
    lines = text.splitlines()
    blocks = []
    current = []
    i = 0

    def flush_paragraph():
        nonlocal current
        if current:
            blocks.append(('para', ' '.join(part.strip() for part in current if part.strip()).strip()))
            current = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if line.startswith('# '):
            flush_paragraph()
            blocks.append(('h1', line[2:].strip()))
            i += 1
            continue
        if line.startswith('## '):
            flush_paragraph()
            blocks.append(('h2', line[3:].strip()))
            i += 1
            continue
        if line.startswith('### '):
            flush_paragraph()
            blocks.append(('h3', line[4:].strip()))
            i += 1
            continue

        image_match = re.fullmatch(r'!\[(.*?)\]\((.*?)\)', stripped)
        if image_match:
            flush_paragraph()
            blocks.append(('image', {'caption': image_match.group(1).strip(), 'path': image_match.group(2).strip()}))
            i += 1
            continue

        if stripped.startswith('|') and stripped.endswith('|'):
            flush_paragraph()
            raw_rows = []
            while i < len(lines):
                row = lines[i].strip()
                if row.startswith('|') and row.endswith('|'):
                    raw_rows.append([cell.strip() for cell in row.strip('|').split('|')])
                    i += 1
                else:
                    break
            if len(raw_rows) >= 2:
                sep = raw_rows[1]
                if all(re.fullmatch(r'[:\- ]+', cell or '-') for cell in sep):
                    raw_rows.pop(1)
            blocks.append(('table', raw_rows))
            continue

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        current.append(line.rstrip())
        i += 1

    flush_paragraph()
    return blocks


def compute_body_word_count(blocks):
    words = []
    in_refs = False
    in_appendix = False
    for kind, value in blocks:
        if kind == 'h2':
            lower = value.strip().lower()
            if lower == 'references':
                in_refs = True
                in_appendix = False
            elif lower.startswith('appendix'):
                in_refs = False
                in_appendix = True
            else:
                in_refs = False
                in_appendix = False
            continue
        if in_refs or in_appendix:
            continue
        if kind == 'h1':
            continue
        if kind == 'para' and value.startswith('Word count'):
            continue
        if kind in {'table', 'image', 'h3'}:
            continue
        if kind == 'para':
            words.append(value)
    body_text = ' '.join(words)
    return len(re.findall(r"\b[\w'-]+\b", body_text))


def extract_reference_records(blocks):
    refs = []
    in_refs = False
    for kind, value in blocks:
        if kind == 'h2':
            lower = value.strip().lower()
            if lower == 'references':
                in_refs = True
                continue
            if lower.startswith('appendix'):
                in_refs = False
        if in_refs and kind == 'para':
            clean = value.replace('*', '')
            match = re.match(r'^(.*?)\s+\((\d{4}[a-z]?)', clean)
            if not match:
                continue
            author_block = match.group(1).strip().rstrip('.')
            year = match.group(2)
            surnames = re.findall(r"([A-Z][A-Za-z'\- ]+),\s*[A-Z]", author_block)
            if surnames:
                if len(surnames) == 1:
                    author_forms = [surnames[0]]
                elif len(surnames) == 2:
                    author_forms = [f'{surnames[0]} & {surnames[1]}', f'{surnames[0]} and {surnames[1]}']
                else:
                    author_forms = [f'{surnames[0]} et al.']
            else:
                author_forms = [author_block]
            refs.append(
                ReferenceRecord(
                    raw_text=value,
                    clean_text=clean,
                    bookmark=sanitize_bookmark(author_block + '_' + year, len(refs) + 1),
                    author_block=author_block,
                    year=year,
                    author_forms=author_forms,
                )
            )
    return refs


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_plain_text(paragraph, text, size=12, bold=False):
    if not text:
        return
    parts = re.split(r'(\*[^*]+\*)', text)
    for part in parts:
        if not part:
            continue
        italic = part.startswith('*') and part.endswith('*') and len(part) >= 2
        literal = part[1:-1] if italic else part
        run = paragraph.add_run(literal)
        set_run_font(run, size=size, bold=bold, italic=italic)


def add_bookmark(paragraph, bookmark_name, counter):
    run = paragraph.add_run('')
    set_run_font(run, size=12)
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(counter))
    start.set(qn('w:name'), bookmark_name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(counter))
    run._r.addprevious(start)
    run._r.addnext(end)


def append_internal_hyperlink(paragraph, text, anchor, size=12):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)
    hyperlink.set(qn('w:history'), '1')

    run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')

    fonts = OxmlElement('w:rFonts')
    fonts.set(qn('w:ascii'), 'Times New Roman')
    fonts.set(qn('w:hAnsi'), 'Times New Roman')
    fonts.set(qn('w:eastAsia'), 'Times New Roman')
    rpr.append(fonts)

    size_el = OxmlElement('w:sz')
    size_el.set(qn('w:val'), str(int(size * 2)))
    rpr.append(size_el)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')
    rpr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'none')
    rpr.append(underline)

    run.append(rpr)
    text_el = OxmlElement('w:t')
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def render_parenthetical_citation(paragraph, content, author_year_lookup, size=12):
    parts = [segment.strip() for segment in content.split(';')]
    parsed = []
    for part in parts:
        match = re.match(r'^(.*?),\s*(\d{4}[a-z]?(?:,\s*\d{4}[a-z]?)*?)$', part)
        if not match:
            return False
        author = match.group(1).strip()
        years = [year.strip() for year in match.group(2).split(',')]
        year_bookmarks = []
        for year in years:
            bookmark = author_year_lookup.get((normalize(author), year))
            if not bookmark:
                return False
            year_bookmarks.append((year, bookmark))
        parsed.append((author, year_bookmarks))

    add_plain_text(paragraph, '(', size=size)
    for seg_index, (author, year_bookmarks) in enumerate(parsed):
        if seg_index:
            add_plain_text(paragraph, '; ', size=size)
        add_plain_text(paragraph, author, size=size)
        add_plain_text(paragraph, ', ', size=size)
        for year_index, (year, bookmark) in enumerate(year_bookmarks):
            if year_index:
                add_plain_text(paragraph, ', ', size=size)
            append_internal_hyperlink(paragraph, year, bookmark, size=size)
    add_plain_text(paragraph, ')', size=size)
    return True


def build_narrative_lookup(reference_records):
    lookup = {}
    for record in reference_records:
        for author_form in record.author_forms:
            lookup[f'{author_form} ({record.year})'] = record.bookmark
    ordered = sorted(lookup.keys(), key=len, reverse=True)
    pattern = re.compile('|'.join(re.escape(item) for item in ordered)) if ordered else None
    return lookup, pattern


def add_text_with_narratives(paragraph, text, narrative_lookup, narrative_pattern, size=12):
    if not text:
        return
    if not narrative_pattern:
        add_plain_text(paragraph, text, size=size)
        return
    pos = 0
    for match in narrative_pattern.finditer(text):
        add_plain_text(paragraph, text[pos:match.start()], size=size)
        literal = match.group(0)
        append_internal_hyperlink(paragraph, literal, narrative_lookup[literal], size=size)
        pos = match.end()
    add_plain_text(paragraph, text[pos:], size=size)


def add_paragraph_with_citations(paragraph, text, author_year_lookup, narrative_lookup, narrative_pattern, size=12):
    pos = 0
    for match in re.finditer(r'\(([^()]*\d{4}[a-z]?[^()]*)\)', text):
        before = text[pos:match.start()]
        add_text_with_narratives(paragraph, before, narrative_lookup, narrative_pattern, size=size)
        content = match.group(1)
        if not render_parenthetical_citation(paragraph, content, author_year_lookup, size=size):
            add_text_with_narratives(paragraph, match.group(0), narrative_lookup, narrative_pattern, size=size)
        pos = match.end()
    add_text_with_narratives(paragraph, text[pos:], narrative_lookup, narrative_pattern, size=size)


def add_table(doc, rows, appendix_mode=False, wide_mode=False):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    if wide_mode and len(rows[0]) == 6:
        widths = [1.05, 0.70, 1.20, 1.25, 1.15, 1.15]
    elif appendix_mode and len(rows[0]) == 5:
        widths = [1.35, 1.15, 2.00, 0.90, 2.10]
    else:
        widths = None

    font_size = 8.5 if wide_mode else (9 if appendix_mode else 10)

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if widths and c_idx < len(widths):
                cell.width = Inches(widths[c_idx])
            cell.text = ''
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            add_plain_text(para, cell_text, size=font_size, bold=(r_idx == 0))
    doc.add_paragraph('')


def main():
    text = SRC.read_text(encoding='utf-8-sig')
    blocks = parse_blocks(text)
    word_count = compute_body_word_count(blocks)
    reference_records = extract_reference_records(blocks)

    author_year_lookup = {}
    for record in reference_records:
        for author_form in record.author_forms:
            author_year_lookup[(normalize(author_form), record.year)] = record.bookmark
    narrative_lookup, narrative_pattern = build_narrative_lookup(reference_records)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)
    section.right_margin = Pt(90)

    styles = doc.styles
    styles['Normal'].font.name = 'Times New Roman'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    styles['Normal'].font.size = Pt(12)
    styles['Normal'].paragraph_format.line_spacing = 1.5
    styles['Normal'].paragraph_format.space_after = Pt(0)

    for style_name in ('Heading 1', 'Heading 2', 'Heading 3'):
        styles[style_name].font.name = 'Times New Roman'
        styles[style_name]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    in_refs = False
    in_appendix = False
    bookmark_counter = 1
    first_appendix = True
    current_h2 = ''
    current_h3 = ''
    ref_iter = iter(reference_records)

    for kind, value in blocks:
        if kind == 'h1':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            set_run_font(run, size=14, bold=True)
            continue

        if kind == 'h2':
            lower = value.strip().lower()
            current_h2 = value.strip()
            current_h3 = ''
            if lower == 'references':
                in_refs = True
                in_appendix = False
            elif lower.startswith('appendix'):
                if first_appendix:
                    doc.add_page_break()
                    first_appendix = False
                in_refs = False
                in_appendix = True
            else:
                in_refs = False
                in_appendix = False
            p = doc.add_paragraph(style='Heading 1')
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(value)
            set_run_font(run, size=12, bold=True)
            continue

        if kind == 'h3':
            current_h3 = value.strip()
            p = doc.add_paragraph(style='Heading 2')
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(value)
            set_run_font(run, size=12, bold=True)
            continue

        if kind == 'image':
            image_path = resolve_path(value['path'], BASE / value['path'])
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run()
            run.add_picture(str(image_path), width=Inches(6.0))
            continue

        if kind == 'table':
            wide_mode = in_appendix and current_h2.startswith('Appendix B') and current_h3 == 'B2. Weighted Decision Matrix'
            add_table(doc, value, appendix_mode=in_appendix, wide_mode=wide_mode)
            continue

        if kind != 'para':
            continue

        paragraph_text = value
        if paragraph_text.startswith('Word count'):
            paragraph_text = f'Word count (main body, excluding references and appendix): {word_count}'

        p = doc.add_paragraph(style='Normal')
        if paragraph_text.startswith('Word count'):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
        elif in_refs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Pt(36)
            p.paragraph_format.first_line_indent = Pt(-36)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(24)

        if in_refs:
            record = next(ref_iter)
            add_bookmark(p, record.bookmark, bookmark_counter)
            bookmark_counter += 1
            add_plain_text(p, record.raw_text, size=12)
        else:
            add_paragraph_with_citations(p, paragraph_text, author_year_lookup, narrative_lookup, narrative_pattern, size=12)

    doc.save(OUT)
    print(f'Wrote {OUT}')
    print(f'WORDCOUNT {word_count}')
    print(f'BOOKMARKS {len(reference_records)}')
    print(f'NARRATIVE_LINK_TARGETS {len(narrative_lookup)}')


if __name__ == '__main__':
    main()

