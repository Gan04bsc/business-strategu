from pathlib import Path
import re
import sys
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn

BASE = Path(r'd:\study\Business Strategy')

def resolve_path(arg, default):
    if arg is None:
        return default
    p = Path(arg)
    return p if p.is_absolute() else BASE / p

src = resolve_path(sys.argv[1] if len(sys.argv) > 1 else None, BASE / 'Assessment_revised.md')
out = resolve_path(sys.argv[2] if len(sys.argv) > 2 else None, src.with_suffix('.docx'))

lines = src.read_text(encoding='utf-8-sig').splitlines()

blocks = []
current = []
i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()
    if line.startswith('# '):
        if current:
            blocks.append(('para', '\n'.join(current).strip()))
            current = []
        blocks.append(('h1', line[2:].strip()))
        i += 1
        continue
    if line.startswith('## '):
        if current:
            blocks.append(('para', '\n'.join(current).strip()))
            current = []
        blocks.append(('h2', line[3:].strip()))
        i += 1
        continue
    if stripped.startswith('|') and stripped.endswith('|'):
        if current:
            blocks.append(('para', '\n'.join(current).strip()))
            current = []
        raw_rows = []
        while i < len(lines):
            s = lines[i].strip()
            if s.startswith('|') and s.endswith('|'):
                cells = [c.strip() for c in s.strip('|').split('|')]
                raw_rows.append(cells)
                i += 1
            else:
                break
        if len(raw_rows) >= 2:
            sep = raw_rows[1]
            if all(re.fullmatch(r'[:\- ]+', cell or '-') for cell in sep):
                raw_rows.pop(1)
        blocks.append(('table', raw_rows))
        continue
    if not stripped:
        if current:
            blocks.append(('para', '\n'.join(current).strip()))
            current = []
        i += 1
        continue
    current.append(line.rstrip())
    i += 1
if current:
    blocks.append(('para', '\n'.join(current).strip()))

exclude = False
body_parts = []
for kind, value in blocks:
    if kind == 'h2' and value.strip().lower() == 'references':
        exclude = True
        continue
    if exclude:
        continue
    if kind == 'h1':
        continue
    if kind == 'para' and value.startswith('Word count'):
        continue
    if kind == 'h2':
        continue
    if kind == 'table':
        continue
    body_parts.append(value)
body_text = '\n'.join(body_parts)
word_count = len(re.findall(r"\b[\w'-]+\b", body_text))

blocks = [
    (kind, f'Word count (main body, excluding references and appendix): {word_count}' if kind == 'para' and value.startswith('Word count') else value)
    for kind, value in blocks
]

doc = Document()
section = doc.sections[0]
section.top_margin = Pt(72)
section.bottom_margin = Pt(72)
section.left_margin = Pt(90)
section.right_margin = Pt(90)

styles = doc.styles
styles['Normal'].font.name = 'Times New Roman'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
styles['Normal'].font.size = Pt(12)
styles['Normal'].paragraph_format.line_spacing = 1.5
styles['Normal'].paragraph_format.space_after = Pt(0)

for style_name in ('Heading 1', 'Heading 2'):
    styles[style_name].font.name = 'Times New Roman'
    styles[style_name]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def add_markdown_runs(paragraph, value):
    parts = re.split(r'(\*[^*]+\*)', value)
    for part in parts:
        if not part:
            continue
        if len(part) >= 2 and part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

in_refs = False
for kind, value in blocks:
    if kind == 'h1':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value)
        run.bold = True
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(14)
    elif kind == 'h2':
        in_refs = value.strip().lower() == 'references'
        p = doc.add_paragraph(style='Heading 1')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(value)
        run.bold = True
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(12)
    elif kind == 'table':
        if not value:
            continue
        table = doc.add_table(rows=len(value), cols=len(value[0]))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for r_idx, row in enumerate(value):
            for c_idx, cell_text in enumerate(row):
                cell = table.cell(r_idx, c_idx)
                cell.text = ''
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.add_run(cell_text)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                run.font.size = Pt(10)
                if r_idx == 0:
                    run.bold = True
        doc.add_paragraph('')
    else:
        p = doc.add_paragraph(style='Normal')
        if value.startswith('Word count'):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if in_refs:
                p.paragraph_format.left_indent = Pt(36)
                p.paragraph_format.first_line_indent = Pt(-36)
            else:
                p.paragraph_format.first_line_indent = Pt(24)
        add_markdown_runs(p, value)

for para in doc.paragraphs:
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        if not run.font.size:
            run.font.size = Pt(12)

doc.save(out)
print(f'Wrote {out}')
print(f'WORDCOUNT {word_count}')
